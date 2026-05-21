"""
src/export/report_exporter.py — PDF/Word Report Exporter
=========================================================
Generates a professionally formatted Word document (.docx) from
the analysis text and charts produced by the agent, then converts
it to PDF using docx2pdf (requires Microsoft Word on Windows).

Called after every analysis when the user confirms export.

Output file:
    reports/analysis/TICKER_report_YYYY-MM-DD.pdf
    reports/analysis/TICKER_report_YYYY-MM-DD.docx (intermediate)

Usage:
    from src.export.report_exporter import export_report
    export_report(tickers, analysis_text, chart_paths)
"""

import logging
from datetime import date
from pathlib import Path

logger = logging.getLogger(__name__)


def _get_chart_paths_for_tickers(tickers: list, charts_dir: Path) -> list:
    """
    Collect all chart PNG files relevant to the given tickers.
    Looks in:
      - reports/charts/TICKER/ for per-ticker charts
      - reports/charts/ root for cross-ticker charts (heatmap, distribution, etc.)
      - reports/charts/compare_TICKER1_TICKER2/ for comparison charts

    Args:
        tickers:    List of ticker symbols
        charts_dir: Base charts directory (reports/charts/)

    Returns:
        List of Path objects for existing chart files, in display order
    """
    found = []

    # Per-ticker charts (price, bollinger, candlestick, financial dashboard)
    for ticker in tickers:
        ticker_dir = charts_dir / ticker
        if ticker_dir.exists():
            for chart in sorted(ticker_dir.glob("*.png")):
                found.append(chart)

    # Cross-ticker charts in root (correlation, distribution, comparative returns)
    for chart in sorted(charts_dir.glob("chart*.png")):
        found.append(chart)

    # Comparison charts folder (compare_AAPL_MSFT/)
    for compare_dir in sorted(charts_dir.glob("compare_*")):
        if compare_dir.is_dir():
            # Only include if the tickers match
            dir_name = compare_dir.name  # e.g. compare_AAPL_MSFT
            parts = dir_name.replace("compare_", "").split("_")
            if any(t in parts for t in tickers):
                for chart in sorted(compare_dir.glob("*.png")):
                    found.append(chart)

    # Remove duplicates while preserving order
    seen = set()
    unique = []
    for p in found:
        if p not in seen:
            seen.add(p)
            unique.append(p)

    return unique


def _parse_sections(analysis_text: str) -> list:
    """
    Split the analysis text into sections based on the ── dividers.
    Each section is a dict with 'title' and 'body' keys.

    Args:
        analysis_text: Full analysis string from run_analysis()

    Returns:
        List of dicts: [{"title": "TREND ANALYSIS", "body": "..."}, ...]
    """
    sections = []
    current_title = "Analysis"
    current_body  = []

    for line in analysis_text.split("\n"):
        # Detect section divider lines like "── NVDA TREND ANALYSIS ──────"
        if line.strip().startswith("──") or line.strip().startswith("=="):
            # Save previous section if it has content
            if current_body:
                body_text = "\n".join(current_body).strip()
                if body_text:
                    sections.append({
                        "title": current_title,
                        "body":  body_text,
                    })
            # Extract new section title from the divider line
            clean = line.replace("─", "").replace("=", "").strip()
            current_title = clean if clean else "Section"
            current_body  = []
        else:
            current_body.append(line)

    # Add the last section
    if current_body:
        body_text = "\n".join(current_body).strip()
        if body_text:
            sections.append({"title": current_title, "body": body_text})

    return sections


def export_report(tickers: list,
                  analysis_text: str,
                  chart_paths: list = None,
                  output_dir: Path = None) -> str:
    """
    Generate a formatted PDF report combining analysis text and charts.

    Process:
      1. Build a Word document (.docx) with cover page, analysis sections,
         and embedded chart images
      2. Convert the .docx to PDF using docx2pdf (requires MS Word on Windows)
      3. Delete the intermediate .docx file
      4. Return the path to the final PDF

    Args:
        tickers:      List of ticker symbols analysed
        analysis_text:Full analysis string from run_analysis()
        chart_paths:  Optional list of chart PNG paths to embed.
                      If None, auto-discovers from reports/charts/
        output_dir:   Where to save the report. Defaults to reports/analysis/

    Returns:
        Path to the saved PDF file as a string, or empty string on failure
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""

    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent.parent))
    from config import REPORTS_ANALYSIS_DIR, REPORTS_CHARTS_DIR

    # Use defaults if not provided
    if output_dir is None:
        output_dir = REPORTS_ANALYSIS_DIR

    # Build output filename with date stamp
    today      = date.today().strftime("%Y-%m-%d")
    label      = "_".join(tickers) if tickers else "report"
    docx_path  = output_dir / f"{label}_report_{today}.docx"
    pdf_path   = output_dir / f"{label}_report_{today}.pdf"

    logger.info(f"  Building report document for {tickers}...")

    # ── Create Word document ──────────────────────────────────────────────────
    doc = Document()

    # Set page margins (narrower for more content space)
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

    # ── Cover page ────────────────────────────────────────────────────────────
    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run  = title_para.add_run("FinAgent")
    title_run.bold      = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark blue

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run  = sub_para.add_run("AI-Powered Financial Analysis Report")
    sub_run.font.size  = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Medium blue

    doc.add_paragraph()  # Spacer

    # Report details table
    details = doc.add_table(rows=3, cols=2)
    details.style = "Table Grid"
    cells = [
        ("Tickers analysed", ", ".join(tickers)),
        ("Report date",      date.today().strftime("%d %B %Y")),
        ("Generated by",     "FinAgent — AI-Powered Financial Data Agent"),
    ]
    for i, (label_text, value_text) in enumerate(cells):
        row = details.rows[i]
        # Label cell (bold, blue background)
        label_cell = row.cells[0]
        label_cell.text = label_text
        label_cell.paragraphs[0].runs[0].bold = True
        label_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        # Value cell
        row.cells[1].text = value_text

    doc.add_page_break()

    # ── Analysis sections ─────────────────────────────────────────────────────
    doc.add_heading("Analysis", level=1)

    sections = _parse_sections(analysis_text)

    if not sections:
        # Fallback: dump the raw text if parsing fails
        doc.add_paragraph(analysis_text)
    else:
        for section_data in sections:
            title = section_data["title"].strip()
            body  = section_data["body"].strip()

            if not body:
                continue

            # Section heading
            if title:
                heading = doc.add_heading(title, level=2)
                heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

            # Section body — split into paragraphs
            for para_text in body.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after  = Pt(6)
                p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ── Charts ────────────────────────────────────────────────────────────────
    # Auto-discover charts if not provided
    if chart_paths is None:
        chart_paths = _get_chart_paths_for_tickers(tickers, REPORTS_CHARTS_DIR)

    # Filter to only existing PNG files
    valid_charts = [Path(p) for p in chart_paths
                    if p and Path(p).exists() and str(p).endswith(".png")]

    if valid_charts:
        doc.add_heading("Charts", level=1)

        for chart_path in valid_charts:
            chart_path = Path(chart_path)

            # Section heading from filename
            # e.g. chart1_price_volume.png → Price Volume
            chart_name = chart_path.stem
            chart_name = chart_name.replace("chart1_", "").replace("chart2_", "") \
                                   .replace("chart3_", "").replace("chart4_", "") \
                                   .replace("chart5_", "").replace("chart6_", "") \
                                   .replace("chart7_", "").replace("chart8_", "") \
                                   .replace("C1_", "").replace("C2_", "") \
                                   .replace("C3_", "").replace("C4_", "") \
                                   .replace("_", " ").title()

            doc.add_heading(chart_name, level=2)

            try:
                # Add chart image — width 6 inches fits within margins
                doc.add_picture(str(chart_path), width=Inches(6.0))
                # Centre the image
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Caption
                caption = doc.add_paragraph(f"Figure: {chart_name}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.size  = Pt(9)
                caption.runs[0].font.italic = True
                caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            except Exception as e:
                logger.warning(f"  Could not embed chart {chart_path.name}: {e}")
                doc.add_paragraph(f"[Chart unavailable: {chart_path.name}]")

            doc.add_paragraph()  # Spacer between charts

    # ── Save Word document ────────────────────────────────────────────────────
    doc.save(str(docx_path))
    logger.info(f"  Word document saved: {docx_path.name}")

    # ── Convert to PDF ────────────────────────────────────────────────────────
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        logger.info(f"  PDF exported: {pdf_path.name}")
        # Remove intermediate .docx file
        docx_path.unlink()
        return str(pdf_path)

    except ImportError:
        logger.warning("  docx2pdf not installed — saving as .docx only.")
        return str(docx_path)

    except Exception as e:
        logger.warning(f"  PDF conversion failed: {e}")
        logger.warning("  Microsoft Word may not be installed. Keeping .docx file.")
        return str(docx_path)